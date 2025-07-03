from typing import List, Dict, Any
import os
import docx

from app.utils.chunking import Document

class DocxLoader:
    """Loader for DOCX documents."""
    
    @staticmethod
    def load(file_path: str) -> Document:
        """Load DOCX file and convert to Document object."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        try:
            # Extract text from DOCX
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Extract metadata
            filename = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            
            # Create metadata
            metadata = {
                "source": file_path,
                "filename": filename,
                "file_size": file_size,
                "file_type": "docx"
            }
            
            return Document(content=text, metadata=metadata)
            
        except Exception as e:
            raise Exception(f"Failed to load DOCX: {str(e)}")
